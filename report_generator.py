from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_word_report(
    filepath,               # folder where the .docx will be saved
    report_date,
    group_name,
    declared_total,
    actual_total,
    tally_match,
    summary,                # list of dicts from build_summary()
    results                 # not used in compact version (kept for compatibility)
):
    """
    Generates a one‑page Word report with a compact summary per agent.
    """
    doc = Document()

    # ---- Company header ----
    heading = doc.add_heading("Claypole Trading Ltd", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Report title ----
    title = doc.add_heading("QUALITY VERIFICATION REPORT", level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Metadata ----
    doc.add_paragraph(f"Date: {report_date}")
    doc.add_paragraph(f"Group: {group_name}")
    doc.add_paragraph(f"Declared Total: {declared_total}")
    doc.add_paragraph(f"Actual Total: {actual_total}")
    doc.add_paragraph(f"Tally Status: {'MATCH' if tally_match else 'MISMATCH'}")
    doc.add_paragraph()  # blank line

    # ---- Compact agent table ----
    # Header row
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'   # clean, modern style
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Agent'
    hdr_cells[1].text = 'Total Submitted'
    hdr_cells[2].text = 'Non‑Quality Serials (last 5)'
    hdr_cells[3].text = 'Not Found Serials (last 5)'

    # Populate rows from summary
    for agent_info in summary:
        agent = agent_info['agent']
        total = agent_info['total']
        non_quality_serials = agent_info.get('non_quality_serials', [])
        not_found_serials = agent_info.get('not_found_serials', [])

        nq_str = ', '.join(non_quality_serials) if non_quality_serials else 'None'
        nf_str = ', '.join(not_found_serials) if not_found_serials else 'None'

        row_cells = table.add_row().cells
        row_cells[0].text = agent
        row_cells[1].text = str(total)
        row_cells[2].text = nq_str
        row_cells[3].text = nf_str

    # ---- Save the document ----
    os.makedirs(filepath, exist_ok=True)
    safe_group = group_name.replace(" ", "_")
    safe_date = report_date.replace("/", "-")
    filename = f"{safe_group}_{safe_date}.docx"
    full_path = os.path.join(filepath, filename)
    doc.save(full_path)
    return full_path